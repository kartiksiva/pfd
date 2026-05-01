import json
import re
import tempfile
from io import BytesIO
from html import escape
from datetime import date
from datetime import datetime, timezone
from pathlib import Path
from typing import Dict, List, Optional
from uuid import UUID

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from app.pdd_template import render_document_markdown


def _render_markdown(document: Dict, sipoc: List[Dict], document_type: str) -> str:
    return render_document_markdown(document=document, sipoc=sipoc, document_type=document_type)


def _render_pdf_from_docx(docx_path: Path, target_path: Path) -> None:
    styles = getSampleStyleSheet()
    heading_1 = ParagraphStyle("DocxHeading1", parent=styles["Heading1"], fontSize=16, leading=20, spaceAfter=6)
    heading_2 = ParagraphStyle("DocxHeading2", parent=styles["Heading2"], fontSize=13, leading=17, spaceBefore=6, spaceAfter=4)
    heading_3 = ParagraphStyle("DocxHeading3", parent=styles["Heading3"], fontSize=11, leading=14, spaceBefore=4, spaceAfter=2)
    body = ParagraphStyle("DocxBody", parent=styles["BodyText"], fontSize=10, leading=13, spaceAfter=3)
    bullet = ParagraphStyle("DocxBullet", parent=body, leftIndent=12, bulletIndent=4)
    table_cell = ParagraphStyle("DocxTableCell", parent=body, fontSize=8.6, leading=10, spaceAfter=0, wordWrap="CJK")
    table_header = ParagraphStyle("DocxTableHeader", parent=table_cell, fontName="Helvetica-Bold")

    pdf = SimpleDocTemplate(
        str(target_path),
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=12 * mm,
        bottomMargin=12 * mm,
    )
    story = []

    doc = Document(str(docx_path))
    for block in doc.element.body:
        if block.tag.endswith("}p"):
            paragraph = next((p for p in doc.paragraphs if p._element is block), None)
            text = (paragraph.text if paragraph else "").strip()
            image_blobs = []
            if paragraph:
                for run in paragraph.runs:
                    for blip in run._element.xpath(".//*[local-name()='blip']"):
                        rid = blip.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                        if rid and rid in doc.part.related_parts:
                            image_blobs.append(doc.part.related_parts[rid].blob)

            if not text and not image_blobs:
                story.append(Spacer(1, 2))
                continue
            style_name = str(getattr(paragraph.style, "name", "") or "")
            if text:
                if style_name.startswith("Heading 1"):
                    story.append(Paragraph(escape(text), heading_1))
                elif style_name.startswith("Heading 2"):
                    story.append(Paragraph(escape(text), heading_2))
                elif style_name.startswith("Heading 3"):
                    story.append(Paragraph(escape(text), heading_3))
                elif "List Bullet" in style_name:
                    story.append(Paragraph(escape(text), bullet, bulletText="•"))
                else:
                    story.append(Paragraph(escape(text), body))
            max_width = A4[0] - (24 * mm)
            for blob in image_blobs:
                try:
                    image = RLImage(BytesIO(blob))
                    if float(image.drawWidth) > max_width and float(image.drawWidth) > 0:
                        scale = max_width / float(image.drawWidth)
                        image.drawWidth = max_width
                        image.drawHeight = float(image.drawHeight) * scale
                    story.append(image)
                    story.append(Spacer(1, 6))
                except Exception:
                    continue
            continue

        if block.tag.endswith("}tbl"):
            table = next((t for t in doc.tables if t._element is block), None)
            if not table:
                continue
            table_rows = []
            for row_idx, row in enumerate(table.rows):
                rendered_row = []
                for cell in row.cells:
                    text = escape(cell.text.strip().replace("\n", " "))
                    style = table_header if row_idx == 0 else table_cell
                    rendered_row.append(Paragraph(text or "-", style))
                table_rows.append(rendered_row)
            if not table_rows:
                continue
            col_width = (A4[0] - (24 * mm)) / len(table_rows[0])
            pdf_table = Table(table_rows, repeatRows=1, colWidths=[col_width] * len(table_rows[0]))
            pdf_table.setStyle(
                TableStyle(
                    [
                        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#94a3b8")),
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e2e8f0")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f172a")),
                        ("FONTSIZE", (0, 0), (-1, -1), 8.8),
                        ("LEADING", (0, 0), (-1, -1), 10),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 4),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                        ("TOPPADDING", (0, 0), (-1, -1), 3),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                    ]
                )
            )
            story.append(pdf_table)
            story.append(Spacer(1, 6))

    if not story:
        story.append(Paragraph("No content available for PDF export.", body))
    pdf.build(story)


def _add_labeled_bullet(doc: Document, label: str, value: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(label).bold = True
    p.add_run(value)


def _safe_add_picture(doc: Document, image_path: Path, width: float = 5.8) -> bool:
    try:
        doc.add_picture(str(image_path), width=Inches(width))
        return True
    except Exception:
        try:
            from PIL import Image

            with Image.open(image_path) as img:
                normalized = img.convert("RGB")
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp_path = Path(tmp.name)
                try:
                    normalized.save(tmp_path, format="PNG")
                    doc.add_picture(str(tmp_path), width=Inches(width))
                    return True
                finally:
                    try:
                        tmp_path.unlink(missing_ok=True)
                    except Exception:
                        pass
        except Exception:
            return False


def _add_table(doc: Document, headers: List[str], rows: List[List[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, h in enumerate(headers):
        table.rows[0].cells[idx].text = h
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value


def _strip_markdown_inline(text: str) -> str:
    return text.replace("**", "").replace("__", "").strip()


def _is_markdown_alignment_row(cells: List[str]) -> bool:
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells if cell.strip())


def _parse_markdown_table_row(line: str) -> List[str]:
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def _resolve_markdown_image_path(raw_path: str) -> Optional[Path]:
    candidate = Path(raw_path.strip())
    if candidate.exists():
        return candidate

    normalized = raw_path.strip()
    mappings = [
        ("/app/uploads/", "uploads/"),
        ("/app/uploads/", "data/uploads/"),
    ]
    for old, new in mappings:
        if normalized.startswith(old):
            remapped = Path(normalized.replace(old, new, 1))
            if remapped.exists():
                return remapped
            remapped_parent = Path("..") / remapped
            if remapped_parent.exists():
                return remapped_parent
    return None


def _render_docx_from_markdown(markdown: str, target_path: Path) -> None:
    doc = Document()
    compact_meta_style = doc.styles.add_style("Compact Meta", WD_STYLE_TYPE.PARAGRAPH)
    compact_meta_style.paragraph_format.space_before = Pt(0)
    compact_meta_style.paragraph_format.space_after = Pt(0)
    compact_meta_style.paragraph_format.line_spacing = 1.0
    meta_prefixes = (
        "**Function:**",
        "**Sub-Function:**",
        "**Document Version:**",
        "**Document Status:**",
        "**Effective Date:**",
    )
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            parsed = [_parse_markdown_table_row(row) for row in block]
            parsed = [row for row in parsed if row and not _is_markdown_alignment_row(row)]
            if parsed:
                headers = [_strip_markdown_inline(cell) for cell in parsed[0]]
                rows = [[_strip_markdown_inline(cell) for cell in row] for row in parsed[1:]]
                _add_table(doc, headers, rows)
            continue

        if stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 6)
            text = _strip_markdown_inline(stripped[level:].strip())
            if text:
                paragraph = doc.add_heading(text, level=level)
                paragraph.paragraph_format.space_before = Pt(6)
                paragraph.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        image_match = re.search(r"!\[(.*?)\]\(([^)]+)\)", stripped)
        if image_match:
            alt_text = image_match.group(1).strip() or "Image"
            image_path = _resolve_markdown_image_path(image_match.group(2).strip())
            if image_path and _safe_add_picture(doc, image_path, width=5.8):
                caption = doc.add_paragraph(alt_text)
                caption.paragraph_format.space_before = Pt(0)
                caption.paragraph_format.space_after = Pt(2)
            else:
                missing = doc.add_paragraph(f"{alt_text}: image not available")
                missing.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        if stripped.startswith(("- ", "* ")):
            doc.add_paragraph(_strip_markdown_inline(stripped[2:].strip()), style="List Bullet")
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        clean = _strip_markdown_inline(stripped)
        if any(stripped.startswith(prefix) for prefix in meta_prefixes):
            paragraph = doc.add_paragraph(clean, style=compact_meta_style)
            paragraph.paragraph_format.space_after = Pt(0)
        else:
            paragraph = doc.add_paragraph(clean)
            paragraph.paragraph_format.space_after = Pt(2)
        i += 1

    doc.save(str(target_path))


def _render_docx_pdd(pdd: Dict, sipoc: List[Dict], target_path: Path) -> None:
    doc = Document()
    steps = pdd.get("steps", []) if isinstance(pdd.get("steps"), list) else []
    systems = pdd.get("systems", []) if isinstance(pdd.get("systems"), list) else []
    business_rules = pdd.get("business_rules", []) if isinstance(pdd.get("business_rules"), list) else []
    exceptions = pdd.get("exceptions", []) if isinstance(pdd.get("exceptions"), list) else []
    preconditions = pdd.get("preconditions", []) if isinstance(pdd.get("preconditions"), list) else []
    metrics = pdd.get("metrics", []) if isinstance(pdd.get("metrics"), list) else []
    risks = pdd.get("risks", []) if isinstance(pdd.get("risks"), list) else []
    triggers = pdd.get("triggers", []) if isinstance(pdd.get("triggers"), list) else []
    outputs = pdd.get("outputs", []) if isinstance(pdd.get("outputs"), list) else []

    process_name = (steps[0].get("title", "") if steps else "").strip() or "Analyzed Process"
    purpose = str(pdd.get("purpose", "")).strip() or "Document current process flow from submitted evidence."
    scope = str(pdd.get("scope", "")).strip() or "Current-state process only."

    doc.add_heading(f"{process_name} - Process Definition Document (PDD)", level=1)

    doc.add_heading("1. Document Control", level=2)
    _add_table(
        doc,
        ["Version", "Date", "Author", "Description"],
        [["1.0", date.today().isoformat(), "PFCD Agent", "Initial Draft (As-Is Process)"]],
    )

    doc.add_heading("2. Process Overview", level=2)
    _add_labeled_bullet(doc, "Process Name: ", process_name)
    _add_labeled_bullet(doc, "Objective: ", purpose)
    _add_labeled_bullet(doc, "Frequency: ", triggers[0] if triggers else "TBD")
    _add_labeled_bullet(doc, "Estimated Volume: ", "TBD")
    _add_labeled_bullet(doc, "Manual Effort: ", "TBD")

    doc.add_heading("3. Scope", level=2)
    doc.add_heading("3.1 In-Scope", level=3)
    doc.add_paragraph(scope, style="List Bullet")
    doc.add_heading("3.2 Out-of-Scope", level=3)
    doc.add_paragraph("Future-state redesign and optimization.", style="List Bullet")

    doc.add_heading("4. Prerequisites & Systems", level=2)
    doc.add_heading("4.1 Prerequisites", level=3)
    if preconditions:
        for item in preconditions:
            doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph("None explicitly stated in source evidence.", style="List Bullet")
    doc.add_heading("4.2 Application Inventory", level=3)
    _add_table(
        doc,
        ["Application", "Version", "Access Method"],
        [[str(app), "Unknown", "As observed in evidence"] for app in systems] or [["Unspecified", "Unknown", "Unknown"]],
    )

    doc.add_heading("5. Detailed Process Steps (As-Is)", level=2)
    _add_table(
        doc,
        ["Step #", "Action", "Role", "System", "Input", "Output"],
        [
            [
                str(step.get("step_no", "")),
                str(step.get("description", "") or step.get("title", "")),
                str(step.get("actor", "")),
                str(step.get("system", "")),
                str(step.get("input", "")),
                str(step.get("output", "")),
            ]
            for step in steps
        ]
        or [["1.1", "Step details unavailable", "Unspecified", "Unspecified", "Unspecified", "Unspecified"]],
    )
    doc.add_heading("Step Screenshots", level=3)
    for step in steps:
        step_no = str(step.get("step_no", ""))
        step_title = str(step.get("title", "") or step.get("description", "") or f"Step {step_no}")
        doc.add_paragraph(f"Step {step_no}: {step_title}")
        if step.get("source_timestamp"):
            doc.add_paragraph(f"Source Timestamp: {step.get('source_timestamp')}", style="List Bullet")
        screenshot = step.get("screenshot") if isinstance(step.get("screenshot"), dict) else {}
        screenshot_path = Path(str(screenshot.get("path", "")).strip()) if screenshot else None
        if screenshot_path and screenshot_path.exists():
            if _safe_add_picture(doc, screenshot_path, width=5.8):
                if screenshot.get("reason"):
                    doc.add_paragraph(f"Frame Reason: {screenshot.get('reason')}", style="List Bullet")
            else:
                doc.add_paragraph("Screenshot: File exists but could not be parsed as an image.", style="List Bullet")
        else:
            doc.add_paragraph("Screenshot: Not available", style="List Bullet")

    doc.add_heading("6. Business Rules & Logic", level=2)
    if business_rules:
        for idx, rule in enumerate(business_rules, start=1):
            _add_labeled_bullet(doc, f"Rule {idx}: ", str(rule))
    else:
        doc.add_paragraph("No explicit business rules identified from source evidence.", style="List Bullet")

    doc.add_heading("7. Exceptions Handling", level=2)
    doc.add_heading("7.1 Business Exceptions", level=3)
    if exceptions:
        for exc in exceptions:
            doc.add_paragraph(str(exc), style="List Bullet")
    else:
        doc.add_paragraph("None captured.", style="List Bullet")
    doc.add_heading("7.2 Technical Exceptions", level=3)
    doc.add_paragraph("Retry, log, and escalate to technical owner if persistent.", style="List Bullet")

    doc.add_heading("8. Inputs & Outputs", level=2)
    _add_labeled_bullet(doc, "Primary Input: ", sipoc[0].get("input", "Unspecified") if sipoc else "Unspecified")
    _add_labeled_bullet(doc, "Primary Output: ", outputs[0] if outputs else (sipoc[-1].get("output", "Unspecified") if sipoc else "Unspecified"))

    doc.add_heading("9. Metrics & Risks", level=2)
    if metrics:
        for metric in metrics:
            _add_labeled_bullet(doc, "Success Metric: ", str(metric))
    else:
        _add_labeled_bullet(doc, "Success Metric: ", "Not explicitly defined.")
    if risks:
        for risk in risks:
            _add_labeled_bullet(doc, "Risk: ", str(risk))
            _add_labeled_bullet(doc, "Mitigation: ", "Define owner, threshold, and contingency action.")
    else:
        _add_labeled_bullet(doc, "Risk: ", "No explicit risks identified.")

    doc.add_heading("10. SIPOC", level=2)
    _add_table(
        doc,
        ["Supplier", "Input", "Process", "Output", "Customer"],
        [
            [
                str(row.get("supplier", "")),
                str(row.get("input", "")),
                str(row.get("process_step", "")),
                str(row.get("output", "")),
                str(row.get("customer", "")),
            ]
            for row in sipoc
        ]
        or [["upstream_supplier", "process input", "unspecified", "unspecified", "downstream_customer"]],
    )

    doc.add_paragraph("Document generated for Process Excellence / Automation Readiness.")
    doc.save(str(target_path))


def _render_docx_sop(document: Dict, sipoc: List[Dict], target_path: Path) -> None:
    doc = Document()
    steps = document.get("steps", []) if isinstance(document.get("steps"), list) else []
    doc_control = document.get("document_control", {}) if isinstance(document.get("document_control"), dict) else {}
    scope = document.get("scope", {}) if isinstance(document.get("scope"), dict) else {}
    quality = document.get("quality_checks", {}) if isinstance(document.get("quality_checks"), dict) else {}
    exceptions = document.get("exception_handling", {}) if isinstance(document.get("exception_handling"), dict) else {}
    controls = document.get("controls_and_compliance", {}) if isinstance(document.get("controls_and_compliance"), dict) else {}
    training = document.get("training_and_kt", {}) if isinstance(document.get("training_and_kt"), dict) else {}

    doc.add_heading("Standard Operating Procedure (SOP)", level=1)
    doc.add_heading("Document Control", level=2)
    _add_table(
        doc,
        ["Field", "Details"],
        [
            ["SOP ID", str(doc_control.get("sop_id", "Needs Review"))],
            ["SOP Title", str(doc_control.get("sop_title", "Needs Review"))],
            ["Process Owner", str(doc_control.get("process_owner", "Needs Review"))],
            ["Department / BU", str(doc_control.get("department", "Needs Review"))],
            ["Effective Date", str(doc_control.get("effective_date", "Needs Review"))],
            ["Review Date", str(doc_control.get("review_date", "Needs Review"))],
            ["Version", str(doc_control.get("version", "1.0"))],
            ["Classification", str(doc_control.get("classification", "Internal"))],
        ],
    )
    doc.add_heading("1. Purpose", level=2)
    doc.add_paragraph(str(document.get("purpose", "Needs Review")))

    doc.add_heading("2. Scope", level=2)
    doc.add_heading("2.1 In-Scope", level=3)
    for item in scope.get("in_scope", []) if isinstance(scope.get("in_scope"), list) else ["Needs Review"]:
        doc.add_paragraph(str(item), style="List Bullet")
    doc.add_heading("2.2 Out-of-Scope", level=3)
    for item in scope.get("out_of_scope", []) if isinstance(scope.get("out_of_scope"), list) else ["Needs Review"]:
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("3. Roles & Responsibilities", level=2)
    _add_table(
        doc,
        ["Role", "Responsibility", "Team / Location"],
        [
            [str(r.get("role", "")), str(r.get("responsibility", "")), str(r.get("team_location", ""))]
            for r in (document.get("roles_and_responsibilities", []) or [])
        ]
        or [["Operator", "Execute process steps", "Needs Review"]],
    )

    doc.add_heading("7. Detailed Process Steps", level=2)
    for idx, step in enumerate(steps, start=1):
        doc.add_heading(f"Step {idx} - {step.get('title', f'Step {idx}')}", level=3)
        _add_labeled_bullet(doc, "Performed By: ", str(step.get("actor", "operator")))
        _add_labeled_bullet(doc, "System Used: ", str(step.get("system", "manual_or_unspecified")))
        _add_labeled_bullet(doc, "Source Timestamp: ", str(step.get("source_timestamp", "Needs Review")))
        _add_labeled_bullet(doc, "Action: ", str(step.get("description", "")))
        _add_labeled_bullet(doc, "Expected Result: ", str(step.get("output", "")))
        screenshot = step.get("screenshot") if isinstance(step.get("screenshot"), dict) else {}
        screenshot_path = Path(str(screenshot.get("path", "")).strip()) if screenshot else None
        if screenshot_path and screenshot_path.exists():
            if _safe_add_picture(doc, screenshot_path, width=5.8):
                if screenshot.get("reason"):
                    _add_labeled_bullet(doc, "Frame Reason: ", str(screenshot.get("reason", "")))
            else:
                doc.add_paragraph("Screenshot: File exists but could not be parsed as an image.", style="List Bullet")
        else:
            doc.add_paragraph("Screenshot: Not available", style="List Bullet")

    doc.add_heading("8. Quality Checks & Validation", level=2)
    _add_table(
        doc,
        ["Check #", "What to Validate", "How to Validate", "Done By", "Frequency"],
        [
            [
                str(c.get("check_id", "")),
                str(c.get("what_to_validate", "")),
                str(c.get("how_to_validate", "")),
                str(c.get("done_by", "")),
                str(c.get("frequency", "")),
            ]
            for c in (quality.get("checks", []) or [])
        ]
        or [["QC-01", "Needs Review", "Needs Review", "Needs Review", "Needs Review"]],
    )

    doc.add_heading("9. Exception Handling", level=2)
    _add_table(
        doc,
        ["Exception ID", "Scenario", "Trigger / Symptom", "Action to Take", "Escalation Path"],
        [
            [
                str(e.get("exception_id", "")),
                str(e.get("scenario", "")),
                str(e.get("trigger_symptom", "")),
                str(e.get("action_to_take", "")),
                str(e.get("escalation_path", "")),
            ]
            for e in (exceptions.get("exception_matrix", []) or [])
        ]
        or [["EXC-01", "Needs Review", "Needs Review", "Needs Review", "Needs Review"]],
    )

    doc.add_heading("10. SLA & Performance Targets", level=2)
    _add_table(
        doc,
        ["KPI", "Definition", "Target", "Measurement Method"],
        [
            [
                str(k.get("kpi", "")),
                str(k.get("definition", "")),
                str(k.get("target", "")),
                str(k.get("measurement_method", "")),
            ]
            for k in (document.get("sla_and_performance_targets", []) or [])
        ]
        or [["Accuracy Rate", "Needs Review", "Needs Review", "Needs Review"]],
    )

    doc.add_heading("13. Controls & Compliance", level=2)
    _add_table(
        doc,
        ["Control", "Type", "Description", "Evidence Required"],
        [
            [str(c.get("control", "")), str(c.get("type", "")), str(c.get("description", "")), str(c.get("evidence_required", ""))]
            for c in (controls.get("controls", []) or [])
        ]
        or [["Needs Review", "Needs Review", "Needs Review", "Needs Review"]],
    )
    doc.add_heading("12. Training & Knowledge Transfer", level=2)
    _add_table(
        doc,
        ["Training Module", "Delivery Mode", "Duration", "Frequency"],
        [
            [
                str(t.get("training_module", "")),
                str(t.get("delivery_mode", "")),
                str(t.get("duration", "")),
                str(t.get("frequency", "")),
            ]
            for t in (training.get("training_requirements", []) or [])
        ]
        or [["Needs Review", "Needs Review", "Needs Review", "Needs Review"]],
    )

    doc.add_heading("SIPOC", level=2)
    _add_table(
        doc,
        ["Supplier", "Input", "Process", "Output", "Customer"],
        [
            [
                str(row.get("supplier", "")),
                str(row.get("input", "")),
                str(row.get("process_step", "")),
                str(row.get("output", "")),
                str(row.get("customer", "")),
            ]
            for row in sipoc
        ]
        or [["upstream_supplier", "process input", "unspecified", "unspecified", "downstream_customer"]],
    )
    doc.save(str(target_path))


def _safe_name_token(raw: str) -> str:
    chars = []
    for ch in (raw or "").strip().lower():
        if ch.isalnum():
            chars.append(ch)
        elif ch in {" ", "-", "_"}:
            chars.append("_")
    token = "".join(chars).strip("_")
    while "__" in token:
        token = token.replace("__", "_")
    return token[:80] or "process"


def _resolve_process_name(pdd: Dict, provided_name: Optional[str]) -> str:
    if provided_name and provided_name.strip():
        return provided_name.strip()
    steps = pdd.get("steps", []) if isinstance(pdd.get("steps"), list) else []
    if steps and str(steps[0].get("title", "")).strip():
        return str(steps[0]["title"]).strip()
    return "process"


def generate_exports(
    job_id: str,
    document: Dict,
    sipoc: List[Dict],
    exports_root: Path,
    document_type: str = "pdd",
    process_name: Optional[str] = None,
    llm_provider: Optional[str] = None,
    processed_at: Optional[datetime] = None,
) -> Dict:
    try:
        UUID(job_id)
    except (ValueError, AttributeError) as exc:
        raise ValueError(f"Invalid job_id format: {job_id!r}") from exc

    job_dir = exports_root / job_id
    artifacts_dir = job_dir / "artifacts"
    artifacts_dir.mkdir(parents=True, exist_ok=True)

    processed_at = processed_at or datetime.now(timezone.utc)
    process_token = _safe_name_token(_resolve_process_name(document, process_name))
    provider_token = _safe_name_token(llm_provider or "provider")
    date_token = processed_at.strftime("%Y%m%d")
    base_name = f"{process_token}_processed_{date_token}_{provider_token}"

    md_content = _render_markdown(document, sipoc, document_type)
    md_path = artifacts_dir / f"{base_name}.md"
    json_path = artifacts_dir / f"{base_name}.json"
    pdf_path = artifacts_dir / f"{base_name}.pdf"
    docx_path = artifacts_dir / f"{base_name}.docx"

    md_path.write_text(md_content, encoding="utf-8")
    json_path.write_text(json.dumps({"document_type": document_type, "document": document, "sipoc": sipoc}, indent=2), encoding="utf-8")
    if document_type == "custom_sop":
        _render_docx_from_markdown(md_content, docx_path)
    elif document_type == "sop":
        _render_docx_sop(document, sipoc, docx_path)
    else:
        _render_docx_pdd(document, sipoc, docx_path)
    _render_pdf_from_docx(docx_path, pdf_path)

    return {
        "md": str(md_path.resolve()),
        "json": str(json_path.resolve()),
        "pdf": str(pdf_path.resolve()),
        "docx": str(docx_path.resolve()),
    }
